"""
converter.py — Universal File Conversion Engine
================================================
Supported conversions:
  PDF  → DOCX, TXT, Images (ZIP)
  DOCX → PDF
  TXT  → PDF
  JPG/PNG/WEBP → PDF (Single image or multi-image compilation)

All public functions return:
  {"ok": True,  "path": "/abs/path/to/output", "filename": "result.xyz"}
  {"ok": False, "error": "human-readable message"}
"""

import io
import os
import re
import uuid
import zipfile
import logging
import tempfile
from pathlib import Path

from PIL import Image

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer


log = logging.getLogger(__name__)

# ── Temporary output directory ─────────────────────────────────────────────
TEMP_DIR = Path(tempfile.gettempdir()) / "student_pdf_toolkit_conversions"
TEMP_DIR.mkdir(parents=True, exist_ok=True)

# ── File-size ceiling: 100 MB ──────────────────────────────────────────────
MAX_BYTES = 100 * 1024 * 1024

# ── MIME / extension maps ──────────────────────────────────────────────────
ALLOWED_INPUT_EXTENSIONS = {".pdf", ".docx", ".txt", ".jpg", ".jpeg", ".png", ".webp"}

OUTPUT_FORMATS: dict[str, list[str]] = {
    ".pdf":  ["docx", "txt", "images"],
    ".docx": ["pdf"],
    ".txt":  ["pdf"],
    ".jpg":  ["pdf"],
    ".jpeg": ["pdf"],
    ".png":  ["pdf"],
    ".webp": ["pdf"],
}

# Type to format mapping
TYPE_TO_FORMAT = {
    "pdf-to-docx":   ("pdf", "docx"),
    "pdf-to-txt":    ("pdf", "txt"),
    "pdf-to-images": ("pdf", "images"),
    "docx-to-pdf":   ("docx", "pdf"),
    "txt-to-pdf":    ("txt", "pdf"),
    "images-to-pdf": ("images", "pdf"),
}


# ══════════════════════════════════════════════════════════════════════════════
#  Public helpers
# ══════════════════════════════════════════════════════════════════════════════

def allowed_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_INPUT_EXTENSIONS


def valid_output_formats(filename: str) -> list[str]:
    ext = Path(filename).suffix.lower()
    return OUTPUT_FORMATS.get(ext, [])


def convert(input_path: str, output_format: str) -> dict:
    """
    Master dispatcher for single-file conversion.
    input_path    — absolute path to the uploaded file (already saved to disk)
    output_format — one of: docx, txt, images, pdf
    """
    src = Path(input_path)
    ext = src.suffix.lower()

    if not src.exists():
        return _err("Uploaded input file not found.")

    # Validate size
    if src.stat().st_size > MAX_BYTES:
        return _err("File exceeds the 100 MB limit.")

    # Validate format combination
    allowed = OUTPUT_FORMATS.get(ext, [])
    if output_format not in allowed:
        return _err(
            f"Cannot convert {ext.lstrip('.')} → {output_format}. "
            f"Allowed targets: {', '.join(allowed) or 'none'}."
        )

    try:
        dispatch = {
            (".pdf",  "docx"):   _pdf_to_docx,
            (".pdf",  "txt"):    _pdf_to_txt,
            (".pdf",  "images"): _pdf_to_images,
            (".docx", "pdf"):    _docx_to_pdf,
            (".txt",  "pdf"):    _txt_to_pdf,
            (".jpg",  "pdf"):    _image_to_pdf,
            (".jpeg", "pdf"):    _image_to_pdf,
            (".png",  "pdf"):    _image_to_pdf,
            (".webp", "pdf"):    _image_to_pdf,
        }
        fn = dispatch.get((ext, output_format))
        if fn is None:
            return _err("Unsupported conversion combination.")
        return fn(src)

    except Exception as exc:
        log.exception("Conversion failed: %s → %s", input_path, output_format)
        return _err(f"Conversion failed: {exc}")


def convert_multiple_images_to_pdf(image_paths: list[str], output_stem: str = "combined_images") -> dict:
    """
    Combine multiple images into a single multi-page PDF.
    image_paths — list of absolute paths to images
    """
    if not image_paths:
        return _err("No images provided for PDF conversion.")

    valid_paths = [Path(p) for p in image_paths if Path(p).exists()]
    if not valid_paths:
        return _err("None of the specified image files exist.")

    try:
        out = _tmp_path(output_stem, ".pdf")
        pil_images = []

        for p in valid_paths:
            try:
                img = Image.open(p)
                # Convert to RGB (handling RGBA, palette, etc.)
                if img.mode in ("RGBA", "LA", "P"):
                    rgb_img = Image.new("RGB", img.size, (255, 255, 255))
                    if img.mode == "RGBA":
                        rgb_img.paste(img, mask=img.split()[3])
                    else:
                        rgb_img.paste(img.convert("RGBA"), mask=img.convert("RGBA").split()[3])
                    pil_images.append(rgb_img)
                else:
                    pil_images.append(img.convert("RGB"))
            except Exception as e:
                log.warning("Could not read image %s: %s", p, e)

        if not pil_images:
            return _err("Failed to decode any valid images.")

        first_img = pil_images[0]
        other_imgs = pil_images[1:] if len(pil_images) > 1 else []

        first_img.save(
            str(out),
            "PDF",
            resolution=100.0,
            save_all=True,
            append_images=other_imgs
        )

        return _ok(out)
    except Exception as exc:
        log.exception("Multiple images to PDF failed")
        return _err(f"Images to PDF conversion failed: {exc}")


def cleanup(path: str):
    """Delete a temporary output file after it has been sent to the client."""
    try:
        p = Path(path)
        if p.exists() and str(TEMP_DIR) in str(p.resolve()):
            p.unlink()
    except OSError:
        pass


# ══════════════════════════════════════════════════════════════════════════════
#  Conversion implementations
# ══════════════════════════════════════════════════════════════════════════════

# ── PDF → DOCX ────────────────────────────────────────────────────────────
def _pdf_to_docx(src: Path) -> dict:
    try:
        from pdf2docx import Converter as Pdf2DocxConverter
    except ImportError:
        return _err("pdf2docx is not installed. Run: pip install pdf2docx")

    out = _tmp_path(src.stem, ".docx")
    cv = Pdf2DocxConverter(str(src))
    try:
        cv.convert(str(out), multi_processing=False)
    finally:
        cv.close()
    return _ok(out)


# ── PDF → TXT (PyMuPDF with PyPDF2 fallback) ──────────────────────────────
def _pdf_to_txt(src: Path) -> dict:
    out = _tmp_path(src.stem, ".txt")
    pages_text = []

    # 1. Try PyMuPDF (fitz) - superior speed and layout extraction
    try:
        import fitz
        doc = fitz.open(str(src))
        for i, page in enumerate(doc, start=1):
            text = page.get_text("text").strip()
            if text:
                pages_text.append(f"--- Page {i} ---\n{text}")
        doc.close()
    except Exception as e:
        log.warning("PyMuPDF text extraction failed: %s, falling back to PyPDF2", e)

    # 2. Fallback to PyPDF2 if fitz didn't extract text
    if not pages_text:
        try:
            import PyPDF2
            with open(src, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                for i, page in enumerate(reader.pages, start=1):
                    content = page.extract_text()
                    if content and content.strip():
                        pages_text.append(f"--- Page {i} ---\n{content.strip()}")
        except Exception as e:
            log.warning("PyPDF2 extraction failed: %s", e)

    if not pages_text:
        return _err(
            "No selectable text found in this PDF. "
            "It may contain scanned images without OCR text."
        )

    out.write_text("\n\n".join(pages_text), encoding="utf-8")
    return _ok(out)


# ── PDF → Images (ZIP of PNGs) ────────────────────────────────────────────
def _pdf_to_images(src: Path) -> dict:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return _err("PyMuPDF is not installed. Run: pip install PyMuPDF")

    out_zip = _tmp_path(src.stem + "_images", ".zip")
    doc = fitz.open(str(src))

    if len(doc) == 0:
        doc.close()
        return _err("The PDF has no pages.")

    with zipfile.ZipFile(out_zip, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, page in enumerate(doc, start=1):
            mat = fitz.Matrix(2.0, 2.0)  # 144 DPI crisp rendering
            pix = page.get_pixmap(matrix=mat)
            buf = io.BytesIO(pix.tobytes("png"))
            zf.writestr(f"page_{i:03d}.png", buf.getvalue())

    doc.close()
    return _ok(out_zip)


# ── DOCX → PDF ────────────────────────────────────────────────────────────
def _docx_to_pdf(src: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        return _err("python-docx is not installed. Run: pip install python-docx")

    doc = Document(str(src))
    out = _tmp_path(src.stem, ".pdf")

    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=11, leading=16, spaceAfter=6,
    )
    h1_style = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontSize=16, leading=22, spaceBefore=14, spaceAfter=6,
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontSize=13, leading=18, spaceBefore=10, spaceAfter=4,
    )

    def _para_style(p):
        s = (p.style.name or "").lower()
        if "heading 1" in s:
            return h1_style
        if "heading 2" in s:
            return h2_style
        return body_style

    elements = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            elements.append(Spacer(1, 6))
            continue
        safe = (text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;"))
        elements.append(Paragraph(safe, _para_style(para)))

    # Also handle simple tables in DOCX
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                safe = (row_text.replace("&", "&amp;")
                                .replace("<", "&lt;")
                                .replace(">", "&gt;"))
                elements.append(Paragraph(f"• {safe}", body_style))
        elements.append(Spacer(1, 6))

    if not elements:
        elements.append(Paragraph("Empty Document", body_style))

    pdf_doc = SimpleDocTemplate(
        str(out),
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )
    pdf_doc.build(elements)
    return _ok(out)


# ── TXT → PDF ─────────────────────────────────────────────────────────────
def _txt_to_pdf(src: Path) -> dict:
    raw = src.read_text(encoding="utf-8", errors="replace")
    out = _tmp_path(src.stem, ".pdf")

    styles = getSampleStyleSheet()
    code_style = ParagraphStyle(
        "TxtBody", parent=styles["Normal"],
        fontName="Helvetica", fontSize=10, leading=15, spaceAfter=3,
    )

    elements = []
    for line in raw.splitlines():
        trimmed = line.strip()
        if not trimmed:
            elements.append(Spacer(1, 6))
            continue
        safe = (trimmed.replace("&", "&amp;")
                       .replace("<", "&lt;")
                       .replace(">", "&gt;"))
        elements.append(Paragraph(safe, code_style))

    if not elements:
        elements.append(Paragraph("Empty Text File", code_style))

    pdf_doc = SimpleDocTemplate(
        str(out), pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )
    pdf_doc.build(elements)
    return _ok(out)


# ── JPG/PNG/WEBP → PDF (Single Image) ─────────────────────────────────────
def _image_to_pdf(src: Path) -> dict:
    return convert_multiple_images_to_pdf([str(src)], output_stem=src.stem)


# ══════════════════════════════════════════════════════════════════════════════
#  Internal utilities
# ══════════════════════════════════════════════════════════════════════════════

def _tmp_path(stem: str, suffix: str) -> Path:
    """Return a unique path inside TEMP_DIR."""
    safe_stem = re.sub(r"[^\w\-]", "_", stem)[:60]
    uid = uuid.uuid4().hex[:8]
    return TEMP_DIR / f"{safe_stem}_{uid}{suffix}"


def _ok(path: Path) -> dict:
    return {"ok": True, "path": str(path), "filename": path.name}


def _err(msg: str) -> dict:
    return {"ok": False, "error": msg}